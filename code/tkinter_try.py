from tkinter import *
import cv2 as cv
from PIL import Image, ImageTk
import numpy as np
import ImageProcessing
import DataManager
from pdf2image import convert_from_path
POPPLER_PATH = "C:\\poppler-20.09.0\\bin"
import docx, os, time

# IMAGE_WIDTH_TO_SHOW = 450
# IMAGE_HEIGTH_TO_SHOW = 600
#
# image_numpy_array = []
# original_numpy_array = []
# points = []
#
#
# def click_CutImage(num_image) :
#     global points, file, img_toCut, top2, w, path_list, image_numpy_array, num_image_to_cut, cv_array, original_numpy_array
#     print (num_image)
#     top2 = Toplevel()
#     top2.title("secend window")
#     points = []
#     w = Canvas(top2, width=500, height=500)
#     cv_array = original_numpy_array[num_image].copy()
#
#     img_toCut = ImageTk.PhotoImage(Image.fromarray(cv_array))
#     image_numpy_array[num_image] = cv_array.copy()
#     num_image_to_cut = num_image
#     w.create_image(0, 0, image=img_toCut, anchor="nw")
#     w.grid(row=0)
#     top2.bind("<Button 1>", CutImage)
#     #print("end cut")
#
# def CutImage(eventorigin):
#     global x, y, points, img_toCut, top2, w, path_list, image_numpy_array, num_image_to_cut, cv_array, root_window
#
#     while(len(points)<4):
#         x = eventorigin.x
#         y = eventorigin.y
#         print(x, y)
#
#         cv_array = cv.circle(cv_array, (x, y), 3, 0, -1)
#         points.append((x, y))
#         if len(points) >= 2:
#             cv_array = cv.line(cv_array, points[-1], points[-2], 50, 3)
#
#         img_toCut = ImageTk.PhotoImage(Image.fromarray(cv_array))
#         w.create_image(0, 0, image=img_toCut, anchor="nw")
#
#         top2.bind("<Button 1>", CutImage)
#         top2.mainloop()
#     print (points)
#     top2.destroy()
#
#     image_numpy_array[num_image_to_cut] = ImageProcessing.WrapImage(image_numpy_array[num_image_to_cut], np.array(points[0:4]))
#     image_list[num_image_to_cut] = ImageTk.PhotoImage(Image.fromarray(image_numpy_array[num_image_to_cut]))
#
#     status = Label(root_window, text = "Image "+str(num_image_to_cut+1)+" of " + str(len(image_list)), bd =1, relief = SUNKEN)
#     my_image = Label(root_window, image = image_list[num_image_to_cut])
#     my_image.grid(row = 0, column = 1, columnspan = 3)
#     button_exit = Button(root_window, text = "Exit",padx = 70, pady = 20, command = exit_program)
#     button_wrap = Button(root_window, text = "try again Cut Image",padx = 70, pady = 20, command = lambda: click_CutImage(num_image_to_cut))
#     if num_image_to_cut == len(path_list)-1:
#         button_next = Button(root_window, text=">>", padx=70, pady=20, state=DISABLED, fg="black")
#     else:
#         button_next = Button(root_window, text=">>", padx=70, pady=20, command=lambda: next(num_image_to_cut + 1),
#                              fg="black")
#
#     if num_image_to_cut ==0:
#         button_previous = Button(root_window, text="<<", padx=70, pady=20, state = DISABLED, fg="black")
#     else:
#         button_previous = Button(root_window, text="<<", padx=70, pady=20, command=lambda: back(num_image_to_cut - 1),
#                                  fg="black")
#
#     button_exit.grid(row = 1, column = 0)
#     button_wrap.grid(row = 1, column = 1,  columnspan = 2, sticky = W+E)
#     button_next.grid(row = 0, column =4)
#     button_previous.grid(row = 0, column = 0)
#     status.grid(row=2, column = 0 , columnspan = 3, sticky = W+E)
#     root_window.mainloop()
#
# def next( image_number):
#     global my_image, button_next, button_previous, status, good_Image, bad_Image, image_list, image_path_list, root_window
#
#     my_image.grid_forget()
#
#     my_image = Label(root_window, image=image_list[image_number])
#
#     if image_number == len(image_list) - 1:
#         button_next = Button(root_window, text=">>", padx=70, pady=20, state=DISABLED, fg="black")
#     else:
#         button_next = Button(root_window, text=">>", padx=70, pady=20, command=lambda: next(image_number + 1),fg="black")
#
#     button_previous = Button(root_window, text="<<", padx=70, pady=20,command=lambda: back(image_number -1), fg="black")
#
#     status = Label(root_window, text="Image " + str(image_number+1) + " of " + str(len(image_list)), bd=1, relief=SUNKEN)
#     button_wrap = Button(root_window, text="Cut Image", padx=70, pady=20, command=lambda: click_CutImage(image_number))
#
#     button_wrap.grid(row=1, column=1, columnspan=2, sticky=W + E)
#     my_image.grid(row=0, column=1, columnspan=3)
#     button_next.grid(row = 0, column =4)
#     button_previous.grid(row=0, column=0)
#     status.grid(row=2, column=0, columnspan=3, sticky = W+E)
#
#
# def back( image_number):
#     global my_image, button_next, button_previous, status, good_Image, bad_Image, image_list, image_path_list, root_window
#     print(image_number)
#     my_image.grid_forget()
#
#     my_image = Label(root_window, image=image_list[image_number])
#     button_next = Button(root_window, text=">>", padx=70, pady=20, command=lambda: next(image_number + 1), fg="black")
#
#     if image_number == 0:
#         button_previous = Button(root_window, text="<<", padx=70, pady=20,state = DISABLED, fg="black")
#     else:
#         button_previous = Button(root_window, text="<<", padx=70, pady=20, command=lambda: back(image_number - 1), fg="black")
#
#     button_wrap = Button(root_window, text="Cut Image", padx=70, pady=20, command=lambda: click_CutImage(image_number))
#     status = Label(root_window, text="Image " + str(image_number +1) + " of " + str(len(image_list)), bd=1, relief=SUNKEN)
#
#     button_wrap.grid(row=1, column=1, columnspan=2, sticky=W + E)
#     my_image.grid(row=0, column=1, columnspan=3)
#     button_next.grid(row = 0, column =4)
#     button_previous.grid(row=0, column = 0)
#     status.grid(row=2, column=0, columnspan=3, sticky = W+E)
#
# def exit_program():
#     global root_window, good_Image, bad_Image
#     root_window.destroy()
#
# def wrap_data(root):
#     global my_image, image_list, status, image_path_list, root_window, path_list, image_numpy_array, root_window, original_numpy_array
#     image_path_list = path_list
#     root_window = root
#     good_Image = []
#     bad_Image = []
#     image_list = []
#     for image_path in path_list:
#         image_array = cv.resize(cv.imread(image_path), (IMAGE_WIDTH_TO_SHOW, IMAGE_HEIGTH_TO_SHOW))
#         image_numpy_array.append(image_array)
#         image_list.append(ImageTk.PhotoImage(Image.fromarray(image_array)))
#     original_numpy_array = image_numpy_array.copy()
#     status = Label(root_window, text = "Image 1 of " + str(len(image_list)), bd =1, relief = SUNKEN)
#
#     my_image = Label(root_window, image = image_list[0])
#     my_image.grid(row = 0, column = 1, columnspan = 3)
#
#     button_exit = Button(root_window, text = "Exit",padx = 70, pady = 20, command = exit_program)
#
#     button_wrap = Button(root_window, text = "Cut Image",padx = 70, pady = 20, command = lambda: click_CutImage(0))
#
#     button_next = Button(root_window, text=">>", padx=70, pady=20, command=lambda: next(1), fg="black")
#     button_previous = Button(root_window, text="<<", padx=70, pady=20, state = DISABLED, fg="black")
#
#     button_exit.grid(row = 1, column = 0)
#     button_wrap.grid(row = 1, column = 1,  columnspan = 2, sticky = W+E)
#     button_next.grid(row = 0, column =4)
#     button_previous.grid(row = 0, column = 0)
#     status.grid(row=2, column = 0 , columnspan = 3, sticky = W+E)
#     root.mainloop()
#
# def new_window_wrap_images(root, path_list_images):
#     global path_list
#     path_list = path_list_images
#     top = Toplevel()
#     top.title("Cropp the image")
#     p = wrap_data(top)
#
#
# # myButton = Button(root, text = "", padx = 50, pady = 20, command = lambda: new_window_wrap_images(root), fg = "black", bg = "green")
# # myButton.grid(row = 1, column = 0)
# class PageCroppImage():
#     def __init__(self, rootTK, path_list_images):
#         self.root = rootTK
#         self.path_list_images = path_list_images
#
#
#
#
#
#
#
#
#
#
#
#
# def main():
#     root = Tk()
#     path_list_images = [img1, img2, img3, img4]
#     new_window_wrap_images(root, path_list_images)
#     root.mainloop()
#
# main()
#
# image = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\Capture1.PNG"
# file = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\Capture1.txt"
#
# boundries = ImageProcessing.GetLineBounds(cv.imread(image,0))
# print(boundries)
# print(len(boundries))
# txt_file = open(file, "r" , encoding="utf-8")
# text = txt_file.read()
# txt_file.close()
# lines = text.split('\n')
# num_lines=0
# for line in lines:
#     if line!="":
#         num_lines+=1
# print(num_lines)
# newImagesForTrain = []
#
# if (num_lines == len(boundries)):
#
#     for i in range(len(boundries)):
#         x, y, w, h =  boundries[i]
#         cutImage = image.cutImage(image.imageArray, x, y, x + w, y + h)
#         Label =lines[i]
#         newImagesForTrain.append(ImageProcessing(cutImage, imagePath = None, Label = Label, handwrite_ID = image.handwrite_ID))

# #split to text file of 6 words
# doc = docx.Document(r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\hebrew_text.docx")
# j=0
# f = open(os.path.join(r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts",  "test_word_file17.txt"), "w+", encoding="utf-8")
# for paragraph in doc.paragraphs:
#     list_words = paragraph.text.split(" ")
#     list_words = [word for word in list_words if word!="" and word!="\n" and word!="\t" and ('\xa0' not in word) and len(word)<18 and len(word)>1]
#     #print(list_words)
#     numword = 0
#     while numword< len(list_words) and len(list_words)>0:
#         line = ""
#         for i in range(6):
#
#             numword += 1
#             if numword < len(list_words):
#                 if i==0:
#                     line = list_words[numword]
#                 else:
#                     line= line+" "+list_words[numword]
#             else:
#                 break
#         if i>1:
#             if line!="" and paragraph.text == doc.paragraphs[-1].text:
#                 f.write(line)
#             else:
#                 f.write(line+"\n")
#         print(list(line))
#         print(len(list_words))
#         print(numword)
#
# f.close
#

# txt_file = open(os.path.join(r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts",  "test_word_file2.txt"), "r",encoding="utf-8")
# text = txt_file.read()
# txt_file.close()
# lines = text.split("\n")
# list_fonts = ['Guttman Yad-Brush', 'Guttman Yad-Light', 'Dana Yad AlefAlefAlef Normal', 'Tamir', 'Ktav Yad CLM', 'OS Luizi Round_FFC'\
#               'Liron', "shmuel", "Ben Gurion", "FtPilKahol 1.00", 'Gveret Levin AlefAlefAlef', "Anka CLM"]
# folder_name = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\docx_fonts"
# for style in list_fonts:
#     document = docx.Document()
#     for line in lines:
#         document.add_paragraph( line , style=style)
#     document.save(os.path.join(folder_name, 'demo.docx'))

#
# outputpath = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\docx_fonts_6\anka_clm_2"
# file = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\docx_fonts_6\docx_Anka_CLM.pdf"
# def ExtractImagesFromPDF(file):
#     images = convert_from_path(file, fmt="jpeg", poppler_path =POPPLER_PATH)
#     namefile = os.path.basename(file)
#     name = os.path.splitext(namefile)[0]
#     i = 0
#     for image in images:
#         # image = Image.open(im)
#         new_path_image = os.path.join(outputpath, name +"_"+ str(i) + ".tif")
#         # j=0
#         # while (new_path_image in files):
#         #     new_path_image = os.path.join(outputpath, handwrite_ID + "_"+str(j)+"_" + str(order[i]) + ".tif")
#         #     j += 1
#         i += 1
#         image.save(new_path_image, 'TIFF')
#     return i

#
#
# folder = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\from_docs_to_image"
# files = os.listdir(folder)
# for file in files:
#     if file[-4:].lower() ==".pdf":
#         outputpath, namefile = os.path.split(file)
#         os.mkdir(os.path.join(folder, namefile[:-4])+"_images_")
#         output_folder = os.path.join(folder, namefile[:-4])+"_images_"
#         pdf_file = os.path.join(folder, file)
#         try:
#             # Wait for 20 seconds
#             images = convert_from_path(pdf_file, fmt="jpeg", poppler_path=POPPLER_PATH)
#
#             namefile = os.path.basename(pdf_file)
#             name = os.path.splitext(namefile)[0]
#             i = 0
#             for image in images:
#                 # image = Image.open(im)
#                 new_path_image = os.path.join(output_folder, name + "_" + str(i) + ".tif")
#                 i += 1
#                 image.save(new_path_image, 'TIFF')
#         except MemoryError as error:
#             print ("error : "+str(i))

#
# def Insert_to_database(images_processed, folderSave):
#     numImage = 0
#     image_path_list = []
#     for imageP in images_processed:
#         nameImage = imageP.handwrite_ID +"_"+ str(numImage)
#         pathNewImage = os.path.join(folderSave, nameImage+".tif")
#         image_path_list.append(pathNewImage)
#         im = Image.fromarray(imageP.imageArray)
#         im.save(pathNewImage, 'TIFF')
#         f = open(os.path.join(folderSave, nameImage+".gt.txt"), "w+", encoding="utf-8")
#         f.write(imageP.Label)
#         f.close()
#         numImage += 1
#
# list_pats = []
# txtfile = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\from_docs_to_image\text.txt"
#
# folder = r"C:\Users\Adi Rosental\Documents\she_code\shecode_final_project\handwriteDoc\trainFonts\from_docs_to_image\docx_tamir_images_"
# #handwrite_ID = "docx_luizi"
# handwrite_ID = os.path.basename(folder.rstrip("_images_"))
# os.mkdir(os.path.join(folder, handwrite_ID))
# folderSave = os.path.join(folder, handwrite_ID)
#
# files = os.listdir(folder)
#
# NUMPIXELS = 20
# till_page = len(files)-2
# print(till_page)
# txt_file = open(txtfile, "r" , encoding="utf-8")
# text = txt_file.read()
# txt_file.close()
# lines_txt = text.split('\n')
# num_lines_txt = len(lines_txt)
# print(num_lines_txt)
# num_line=0
# newImagesForTrain = []
# num_lines_images = 0
#
# for num_page in range(till_page):
#     image = os.path.join(folder, handwrite_ID+"_"+str(num_page)+".tif")
#     print(image)
#     imageArray = cv.imread(image, 0)
#     width = imageArray.shape[1]
#     height = imageArray.shape[0]
#     boundries = ImageProcessing.GetLineBounds(imageArray)
#     print(boundries)
#     print(len(boundries))
#     num_lines_images += len(boundries)
#     for i in range(len(boundries)):
#         x, y, w, h = boundries[i]
#         cutImage = imageArray[max(0, min(y, y+h)-NUMPIXELS):min(max(y, y+h)+NUMPIXELS,height) , max(0, min(x, x+w)-NUMPIXELS) :min(width, max(x, x+w)+NUMPIXELS)]
#         while lines_txt[num_line]=="" or lines_txt[num_line]==" ":
#             num_line+=1
#         Label = lines_txt[num_line]
#         num_line += 1
#         # print(Label)
#         # print(list(Label))
#         # cv.imshow("cutImage", cutImage)
#         # cv.waitKey()
#         # print(num_page)
#
#         newImagesForTrain.append(ImageProcessing.ImageProcessing(cutImage, imagePath=None, Label=Label, handwrite_ID=handwrite_ID))
#
# Insert_to_database(newImagesForTrain, folderSave)
# print(num_line, num_lines_images)
# # # return (numS, numE)

folder = r"C:\Users\Adi Rosental\Documents\dato_to_pass\DataBase"
files = os.listdir(folder)
for file in files:
    #outputpath, namefile = os.path.split(file)
    os.rename(os.path.join(folder,file), os.path.join(folder, file[5:]))